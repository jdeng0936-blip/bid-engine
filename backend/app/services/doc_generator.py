import asyncio
"""

流程：
  1. 加载 ProjectParams + Project 基础信息
  2. 调用 RuleService.match_rules() → 命中规则+章节列表
  3. 调用 SupportCalcEngine + VentCalcEngine → 计算结果
  4. 按章节顺序组装内容
  5. python-docx 生成 .docx 文件
"""
import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.project import Project, ProjectParams
from app.schemas.calc import SupportCalcInput
from app.schemas.vent import VentCalcInput
from app.schemas.doc import ChapterContent, DocGenerateResult
from app.services.calc_engine import SupportCalcEngine
from app.services.vent_engine import VentCalcEngine
from app.services.rule_service import RuleService
from app.core.prompt_manager import prompt_manager
# 参数字段中文映射
PARAM_LABELS: dict[str, str] = {
    "rock_class": "围岩级别", "coal_thickness": "煤层厚度(m)",
    "coal_dip_angle": "煤层倾角(°)", "gas_level": "瓦斯等级",
    "hydro_type": "水文地质类型", "geo_structure": "地质构造",
    "spontaneous_combustion": "自燃倾向性", "roadway_type": "巷道类型",
    "excavation_type": "掘进类型", "section_form": "断面形式",
    "section_width": "断面宽度(m)", "section_height": "断面高度(m)",
    "excavation_length": "掘进长度(m)", "service_years": "服务年限(年)",
    "dig_method": "掘进方式", "dig_equipment": "掘进设备",
    "transport_method": "运输方式",
}


class DocGenerator:
    """文档生成引擎"""

    def __init__(self, session: AsyncSession):
        self.session = session

    async def generate(
        self, project_id: int, tenant_id: int, include_calc: bool = True
    ) -> DocGenerateResult:
        """
        端到端文档生成

        Returns:
            DocGenerateResult 包含文件路径和章节列表
        """
        # 1. 加载项目信息
        project = await self._load_project(project_id, tenant_id)
        if not project:
            raise ValueError("项目不存在或无权访问")

        params = await self._load_params(project_id)
        params_dict = self._params_to_dict(params) if params else {}

        # 2. 规则匹配
        rule_svc = RuleService(self.session)
        match_result = None
        try:
            match_result = await rule_svc.match_rules(project_id, tenant_id)
        except ValueError:
            pass  # 参数未填写时跳过匹配

        # 3. 计算引擎
        calc_result = None
        vent_result = None
        if include_calc and params:
            calc_result = self._run_support_calc(params_dict)
            vent_result = self._run_vent_calc(params_dict)

        # 3.5 设备材料匹配
        equipment_result = None
        try:
            from app.services.equipment_material_engine import EquipmentMaterialEngine
            eq_engine = EquipmentMaterialEngine(self.session)
            equipment_result = await eq_engine.run_full_match(project_id, tenant_id)
            print(f"🔧 设备材料匹配完成: {equipment_result.total_equipment_count} 台设备, {equipment_result.total_material_types} 种材料")
        except Exception as e:
            print(f"⚠️ 设备材料匹配失败（降级跳过）: {e}")

        # 4. 组装章节
        chapters = self._assemble_chapters(
            project, params_dict, match_result, calc_result, vent_result,
            equipment_result=equipment_result,
        )

        # 4.5 智能深度润色（AI赋能 + 范文 Few-shot + 计算推导注入）
        chapters = await self._ai_polish_content(
            chapters, params_dict, calc_result, vent_result,
            tenant_id=tenant_id, project=project
        )

        # 5. 生成 Word
        file_path = self._render_docx(
            project, chapters, calc_result, vent_result,
            equipment_result=equipment_result,
        )

        total_warnings = 0
        if calc_result:
            total_warnings += len(calc_result.warnings)
        if vent_result:
            total_warnings += len(vent_result.warnings)

        return DocGenerateResult(
            project_id=project_id,
            project_name=project.face_name,
            file_path=file_path,
            total_chapters=len(chapters),
            total_warnings=total_warnings,
            chapters=chapters,
        )

    # ========== 加载数据 ==========

    async def _load_project(self, pid: int, tid: int) -> Optional[Project]:
        result = await self.session.execute(
            select(Project).where(Project.id == pid, Project.tenant_id == tid)
        )
        return result.scalar_one_or_none()

    async def _load_params(self, pid: int) -> Optional[ProjectParams]:
        result = await self.session.execute(
            select(ProjectParams).where(ProjectParams.project_id == pid)
        )
        return result.scalar_one_or_none()

    @staticmethod
    def _params_to_dict(params: ProjectParams) -> dict:
        return {
            c.key: getattr(params, c.key)
            for c in params.__table__.columns
            if c.key not in ("id", "project_id")
        }

    # ========== 计算引擎调用 ==========

    @staticmethod
    def _run_support_calc(p: dict):
        try:
            inp = SupportCalcInput(
                rock_class=p.get("rock_class", "III"),
                section_form=p.get("section_form", "矩形"),
                section_width=float(p.get("section_width", 4.5)),
                section_height=float(p.get("section_height", 3.2)),
            )
            return SupportCalcEngine.calculate(inp)
        except Exception:
            return None

    @staticmethod
    def _run_vent_calc(p: dict):
        try:
            inp = VentCalcInput(
                gas_emission=float(p.get("gas_emission", 2.0) or 2.0),
                gas_level=p.get("gas_level", "低瓦斯"),
                section_area=float(p.get("section_width", 4.5)) * float(p.get("section_height", 3.2)),
                excavation_length=float(p.get("excavation_length", 300)),
            )
            return VentCalcEngine.calculate(inp)
        except Exception:
            return None

    # ========== 章节组装 ==========

    def _assemble_chapters(self, project, params_dict, match_result, calc_result, vent_result, equipment_result=None):
        """
        严格按华阳集团《采掘运技术管理规定》附件2
        《掘进工作面作业规程编制内容及提纲》组装规程

        官方9章结构:
          第一章 概述
          第二章 地面相对位置及水文地质概况 (4节)
          第三章 巷道布置及支护说明 (5节)
          第四章 施工工艺 (3节)
          第五章 生产系统 (7节)
          第六章 劳动组织及主要技术经济指标 (2节)
          第七章 煤质管理
          第八章 安全技术措施 (9节)
          第九章 安全风险管控及应急避险
          附录
        """
        chapters: list[ChapterContent] = []

        mine_name = getattr(project, 'mine_name', '—')
        roadway_type = params_dict.get('roadway_type', '—')
        gas_level = params_dict.get('gas_level', '—')
        excavation_type = params_dict.get('excavation_type', '—')

        # ================================================================
        #  第一章  概述
        # ================================================================
        ch1_lines = [
            f"一、巷道名称：{project.face_name}",
            f"  矿井名称：{mine_name}",
            f"  巷道用途/性质：{roadway_type}",
            f"  设计长度：{params_dict.get('excavation_length', '—')} m",
            f"  坡度：按设计确定",
            "",
            "二、特殊技术要求及需要重点说明的问题",
            f"  掘进类型：{excavation_type}",
            f"  瓦斯等级：{gas_level}",
            f"  自燃倾向性：{params_dict.get('spontaneous_combustion', '—')}",
            f"  掘进方式：{params_dict.get('dig_method', '—')}",
            "",
            "三、巷道布置平面图",
            "  （详见附图）",
        ]
        chapters.append(ChapterContent(
            chapter_no="第一章", title="概述",
            content="\n".join(ch1_lines), source="template",
        ))

        # ================================================================
        #  第二章  地面相对位置及水文地质概况（4节）
        # ================================================================
        ch2_lines = [
            "第一节  地面相对位置及邻近采区开采情况",
            "",
            f"一、巷道相应的地面位置、标高，区域内的水体和建筑物对工程的影响",
            f"  （根据矿井实际填写）",
            "",
            "二、巷道与相邻煤（岩）层、邻近巷道的层间关系",
            "  附近已有的采掘情况对工程的影响。",
            "",
            "三、老空区的水、火、瓦斯等对工程的影响分析",
            "",
            "第二节  煤（岩）层赋存特征",
            "",
            "一、煤（岩）层产状、厚度、结构",
            f"  煤层厚度：{params_dict.get('coal_thickness', '—')} m",
            f"  煤层倾角：{params_dict.get('coal_dip_angle', '—')}°",
            f"  坚固性系数(f)：根据实测确定",
            "",
            "二、预测巷道瓦斯涌出量、煤层自然发火倾向等",
            f"  瓦斯等级：{gas_level}",
            f"  自燃倾向性：{params_dict.get('spontaneous_combustion', '—')}",
            "",
            "三、其他煤（岩）层技术特征分析",
            "",
            "四、地层综合柱状图（详见附图）",
            "",
            "五、巷道围岩类别",
            f"  围岩级别：{params_dict.get('rock_class', '—')}",
            "",
            "第三节  地质构造",
            "",
            f"一、巷道煤（岩）层产状，断层、褶曲等地质构造要素",
            f"  地质构造特征：{params_dict.get('geo_structure', '—')}",
            "",
            "二、地质平面图、剖面图（详见附图）",
            "",
            "第四节  水文地质",
            "",
            "一、主要充水因素分析",
            f"  水文地质类型：{params_dict.get('hydro_type', '—')}",
            "",
            "二、带压掘进工作面突水系数计算及危险性评价",
            "",
            "三、积水区域附近掘进巷道，标出\"三线\"（积水线、探水线和警戒线）",
            "",
            "四、预测工作面正常、最大涌水量",
        ]
        chapters.append(ChapterContent(
            chapter_no="第二章", title="地面相对位置及水文地质概况",
            content="\n".join(ch2_lines), source="template",
        ))

        # ================================================================
        #  第三章  巷道布置及支护说明（5节）
        # ================================================================
        sec_w = float(params_dict.get("section_width", 0) or 0)
        sec_h = float(params_dict.get("section_height", 0) or 0)
        sec_area = round(sec_w * sec_h, 2) if sec_w > 0 and sec_h > 0 else 0

        ch3_lines = [
            "第一节  巷道布置",
            "",
            f"一、巷道布置：层位、水平标高、开口的位置、方位角",
            f"  巷道类型：{roadway_type}",
            f"  掘进长度：{params_dict.get('excavation_length', '—')} m",
            f"  服务年限：{params_dict.get('service_years', '—')} 年",
            "",
            "二、特殊地点的施工（车场、硐室、交岔点等）",
            "",
            "三、开口大样图（详见附图）",
            "",
            "第二节  矿压观测",
            "",
            "一、采用锚网支护掘进巷道必须安设顶板离层仪和锚杆锚索测力计。",
            "",
            "二、矿压观测分综合监测和一般监测两种：",
            "  综合监测：巷道表面位移、顶板离层、锚杆锚索受力状况监测。",
            "  一般监测：巷道表面位移、顶板离层监测。",
            "",
            "三、回采顺槽巷道设置综合测站和一般测站，明确布站间距、数量及观测分析标准。",
            "",
            "四、特殊地段（开口处、交岔点、构造影响区等）须增设监测仪器。",
            "",
            "第三节  顶板岩性探测",
            "",
            "一、掘进巷道必须进行顶板岩性探测与分析，验证与优化支护设计。",
            "",
            "二、明确岩性探测方法和具体技术要求。",
            "",
            "第四节  支护设计",
            "",
            f"一、巷道断面设计",
            f"  断面形式：{params_dict.get('section_form', '—')}",
            f"  断面宽度：{sec_w} m" if sec_w > 0 else "  断面宽度：—",
            f"  断面高度：{sec_h} m" if sec_h > 0 else "  断面高度：—",
            f"  断面净面积：{sec_area} m²" if sec_area > 0 else "  断面净面积：—",
            "",
        ]

        # 嵌入支护计算结果
        if calc_result:
            ch3_lines.extend([
                "二、支护参数（计算引擎输出）",
                f"  单根锚杆锚固力：{calc_result.bolt_force} kN",
                f"  最大允许锚杆间距：{calc_result.max_bolt_spacing} mm",
                f"  最大允许排距：{calc_result.max_bolt_row_spacing} mm",
                f"  推荐每排锚杆数：{calc_result.recommended_bolt_count_per_row} 根",
                f"  最少锚索数量：{calc_result.min_cable_count} 根",
                f"  支护密度：{calc_result.support_density} 根/m²",
                f"  安全系数：{calc_result.safety_factor}",
                "",
            ])
            if calc_result.warnings:
                ch3_lines.append("【支护合规预警】")
                for w in calc_result.warnings:
                    ch3_lines.append(f"  ⚠ {w.message}")
                ch3_lines.append("")
        else:
            ch3_lines.extend([
                "二、支护参数",
                "  （待填写支护设计参数）",
                "",
            ])

        ch3_lines.extend([
            "三、支护参数校核",
            "  （一）顶锚杆校核：L ≥ L1 + L2 + L3",
            "  （二）校核顶锚杆间排距",
            "  （三）加强锚索长度校核",
            "  （四）加强锚索数目校核",
            "",
            "四、巷道断面图、平面图、交岔点支护示意图（详见附图）",
            "",
            "五、支护设计采用动态信息设计方法",
            "  工程类比法初始设计→持续矿压观测→修改优化→正式设计。",
            "  巷道条件发生变化时，须立即组织现场查看并及时修改支护设计。",
            "",
            "第五节  支护工艺",
            "",
            "一、临时支护工艺、工序及要求",
            "  煤巷综掘、大断面岩巷综掘必须使用机载式临时支护装置。",
            "  临时支护与永久支护距掘进工作面的距离须在规程中明确。",
            "",
            "二、永久支护工艺、工序及要求",
            "  （一）锚杆及联合支护：材质、规格、间排距、安装、锚固力要求。",
            "  （二）支架支护：构件齐全，背紧背牢、充满填实。",
            "",
            "三、施工质量标准表（详见附表）",
        ])

        chapters.append(ChapterContent(
            chapter_no="第三章", title="巷道布置及支护说明",
            content="\n".join(ch3_lines), source="calc_engine" if calc_result else "template",
            has_warning=bool(calc_result and not calc_result.is_compliant),
        ))

        # ================================================================
        #  第四章  施工工艺（3节）
        # ================================================================
        ch4_lines = [
            "第一节  施工方法",
            "",
            "一、确定巷道施工方法",
            f"  掘进方式：{params_dict.get('dig_method', '—')}",
            f"  掘进类型：{excavation_type}",
            "",
            "二、巷道开口施工方法",
            "  从支设临时支护开始，到永久支护止的施工顺序。",
            "",
            "三、特殊条件下的施工方法",
            "  （一）石门揭开煤层时：打超前钻排放瓦斯、远距离放炮。",
            "  （二）硐室的施工方法：根据围岩类别选用全断面或分层施工法。",
            "  （三）交岔点的施工方法：根据围岩类别选用相应施工法。",
            "  （四）倾斜巷道：支架迎山角、防滑防跑车装置。",
            "",
            "第二节  掘进方式",
            "",
            f"一、掘进方式：{params_dict.get('dig_method', '—')}",
            f"  掘进设备：{params_dict.get('dig_equipment', '—')}",
            "",
            "二、机掘作业方式、截割顺序、截割循环进度",
            "",
            "三、炮掘施工工序安排、工艺流程",
            "  严格执行\"一炮三检\"和\"三人连锁放炮\"制度。",
            "  炮孔内发现异状、温度骤高骤低、有显著瓦斯涌出、煤岩松软时，",
            "  必须停止装药，并采取安全措施。",
            "",
            "第三节  装载运输",
            "",
            f"一、运输方式：{params_dict.get('transport_method', '—')}",
            "",
            "二、装载设备及操作要求",
            "",
            "三、管线及轨道敷设",
            "  风、水管路要同侧敷设，不得与瓦斯抽采管路同侧布置。",
        ]

        chapters.append(ChapterContent(
            chapter_no="第四章", title="施工工艺",
            content="\n".join(ch4_lines), source="template",
        ))

        # ================================================================
        #  第五章  生产系统（7节）
        # ================================================================
        ch5_lines = [
            "第一节  一通三防",
            "",
            "一、通风系统",
        ]
        if vent_result:
            ch5_lines.extend([
                f"  瓦斯涌出法需风量(Q_gas)：{vent_result.q_gas} m³/min",
                f"  人数法需风量(Q_people)：{vent_result.q_people} m³/min",
                f"  炸药法需风量(Q_explosive)：{vent_result.q_explosive} m³/min",
                f"  最终配风量(Q_required)：{vent_result.q_required} m³/min",
                f"  推荐局扇型号：{vent_result.recommended_fan}（{vent_result.fan_power} kW）",
            ])
            if vent_result.warnings:
                ch5_lines.append("  【通风合规预警】")
                for w in vent_result.warnings:
                    ch5_lines.append(f"    ⚠ {w.message}")
        else:
            ch5_lines.append("  （待计算通风参数）")

        ch5_lines.extend([
            "",
            "二、综合防尘",
            "  采用湿式打眼、喷雾降尘、通风除尘等综合防尘措施。",
            "  掘进工作面应安设净化水幕、转载点喷雾装置。",
            "",
            "三、防灭火",
            "  巷道布置及通风设施设置须符合防灭火要求。",
            "  厚煤层工作面进、回风巷开口须按规定设置防灭火设施。",
            "",
            "第二节  压风",
            "",
            "  压风自救装置安设间距不超过200米。",
            "  压风管路规格、连接方式及维护要求。",
            "",
            "第三节  动力（供电）",
            "",
            "  掘进工作面供电须采用三专线路（专用变压器、开关、电缆）。",
            "  电气设备选型及防爆要求。",
            "",
            "第四节  排水",
            "",
            "  掘进工作面须配备排水设备，排水能力须满足最大涌水量要求。",
            f"  水文地质类型：{params_dict.get('hydro_type', '—')}",
            "",
            "第五节  运输",
            "",
            f"  运输方式：{params_dict.get('transport_method', '—')}",
            "  运输设备选型、安全措施。",
            "",
            "第六节  通讯照明",
            "",
            "  掘进工作面应设置有线调度电话和应急通讯设备。",
            "  照明灯具选型及安设要求。",
            "",
            "第七节  供水施救",
            "",
            "  供水施救管路安装要求。",
            "  每隔一定距离设置三通阀门。",
        ])

        chapters.append(ChapterContent(
            chapter_no="第五章", title="生产系统",
            content="\n".join(ch5_lines), source="calc_engine" if vent_result else "template",
            has_warning=bool(vent_result and not vent_result.is_compliant),
        ))

        # ================================================================
        #  第六章  劳动组织及主要技术经济指标（2节）
        # ================================================================
        ch6_lines = [
            "第一节  劳动组织",
            "",
            "一、劳动组织",
            "  实行\"三八\"作业制，两班生产、一班检修。",
            "",
            "二、正规循环作业",
            "  按照\"正规循环作业\"要求组织生产，正规循环率不低于80%。",
            "",
            "三、循环作业图表（详见附表）",
            "",
            "第二节  主要技术经济指标",
            "",
            f"  巷道掘进长度：{params_dict.get('excavation_length', '—')} m",
            f"  断面形式：{params_dict.get('section_form', '—')}",
            f"  断面净面积：{sec_area} m²" if sec_area > 0 else "  断面净面积：—",
            "  循环进度、月进度、工效等（详见技术经济指标表）",
        ]

        chapters.append(ChapterContent(
            chapter_no="第六章", title="劳动组织及主要技术经济指标",
            content="\n".join(ch6_lines), source="template",
        ))

        # ================================================================
        #  第七章  煤质管理（官方新增章节）
        # ================================================================
        ch7_lines = [
            "一、煤质指标",
            "  简要说明煤质指标：灰分、水分、发热量、硫分等。",
            "",
            "二、提高煤质及采出率的措施",
            "  （一）严格控制混矸率，掘进中分层排矸。",
            "  （二）煤岩分装分运，防止煤矸混装。",
            "  （三）减少煤尘飞扬损失。",
            "  （四）合理确定巷道断面，提高煤炭采出率。",
        ]

        chapters.append(ChapterContent(
            chapter_no="第七章", title="煤质管理",
            content="\n".join(ch7_lines), source="template",
        ))

        # ================================================================
        #  第八章  安全技术措施 — 拆分为 9 个独立子章节独立 AI 扩写
        #  范文此章 66,567 字/820 条，是全文最大章节
        # ================================================================

        # 第八章-第一节 一般规定
        ch8s1_lines = [
            "第一节  一般规定",
            "",
            "一、工作面安全管理",
            "  （一）明确工作面安全生产责任制，区队长、班组长、安全员、瓦检员等各岗位安全职责。",
            "  （二）开工前必须召开班前会，传达安全注意事项和当班任务。",
            "  （三）严格执行交接班制度，交接班时必须在现场交接。",
            "  （四）严禁违章指挥、违章作业、违反劳动纪律。",
            "",
            "二、交接班管理",
            "  （一）交接班时必须检查的内容：",
            "    1. 工作面顶板、煤壁状况",
            "    2. 瓦斯浓度、通风设施完好情况",
            "    3. 机电设备运转状况",
            "    4. 安全设施、消防器材完好情况",
            "    5. 文明生产情况",
            "  （二）交接班记录必须由交接双方签字确认。",
            "",
            "三、入井须知",
            "  （一）入井人员必须随身携带自救器，严禁携带烟草和点火物品。",
            "  （二）入井前必须参加班前会，了解当天安全注意事项。",
            "  （三）严格执行出入井检身制度和人员定位管理。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第一节", title="安全技术措施——一般规定",
            content="\n".join(ch8s1_lines), source="template",
        ))

        # 第八章-第二节 顶板管理
        ch8s2_lines = [
            "第二节  顶板管理",
            "",
            "一、掘进工作面顶板管理一般规定",
            f"  （一）围岩级别：{params_dict.get('rock_class', '—')}",
            f"  （二）断面形式：{params_dict.get('section_form', '—')}，宽×高：{sec_w}×{sec_h}m",
            "  （三）掘进工作面必须严格执行敲帮问顶制度，每班开工前由班组长或有经验矿工进行全面检查。",
            "  （四）严禁空顶作业，空顶距离不得超过作业规程规定（≤300mm）。",
            "",
            "二、临时支护管理",
            "  （一）煤巷综掘掘进面必须使用机载式临时支护装置。",
            "  （二）临时支护距掘进工作面的距离不得超过规定值。",
            "  （三）临时支护必须满足初始支护阻力要求。",
            "",
            "三、永久支护质量管理",
            "  （一）锚杆安装扭矩不得低于100N·m，预紧力不得低于60kN。",
            "  （二）锚杆锚固力不得低于设计锚固力（≥70kN）。",
            "  （三）锚杆外露长度10~50mm，托盘紧贴岩面。",
            "  （四）锚索张拉预紧力≥120kN，外露长度150~250mm。",
            "",
            "四、矿压监测",
            "  （一）安装顶板离层仪，安装间距不大于50m。",
            "  （二）安装锚杆锚索测力计，每组3个测站。",
            "  （三）巷道变形量超过预警值时必须采取加强支护措施。",
            "",
            "五、构造段顶板管理",
            "  （一）过断层、褶曲等构造段时必须制定专项措施。",
            "  （二）构造段必须缩小锚杆间排距，增加锚索数量。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第二节", title="安全技术措施——顶板管理",
            content="\n".join(ch8s2_lines), source="template",
        ))

        # 第八章-第三节 一通三防
        ch8s3_lines = [
            "第三节  一通三防",
            "",
            f"一、本工作面瓦斯等级：{gas_level}",
            "",
            "二、通风管理",
            "  （一）局部通风机必须实现双风机双电源自动切换，切换时间≤10秒。",
            "  （二）风筒口距掘进工作面的距离：煤巷≤5m，岩巷≤10m。",
            "  （三）工作面风速：最低不得低于0.25m/s，最高不得超过4.0m/s。",
            "  （四）风筒末端100mm水柱风量满足配风要求。",
            "",
            "三、瓦斯管理",
            "  （一）瓦斯检查频次：",
            f"    {'每班至少检查3次（突出工作面有专人连续监测）' if gas_level == '突出' else '每班至少检查3次'}",
            "  （二）瓦斯超限处理：",
            "    1. 掘进工作面CH₄≥0.5%——停止用电钻打眼",
            "    2. 掘进工作面CH₄≥1.0%——停止工作，切断电源，撤出人员",
            "    3. 掘进工作面CH₄≥1.5%——断电撤人，进行处理",
            "    4. 回风流中CH₄≥0.5%——停止工作，切断电源",
            "",
            "四、综合防尘",
            "  （一）掘进工作面必须采用湿式打眼，使用内喷雾截齿。",
            "  （二）设置净化水幕：距掘进工作面50m以内安设净化水幕，喷雾压力≥1.5MPa。",
            "  （三）转载点喷雾：每个转载点均须安设自动喷雾降尘装置。",
            "  （四）个体防护：所有井下作业人员必须佩戴防尘口罩。",
            "",
            "五、防灭火",
            "  （一）掘进工作面进、回风巷每隔100m设置一组灭火器。",
            f"  （二）自燃倾向性：{params_dict.get('spontaneous_combustion', '—')}。",
            "  （三）电气设备着火时必须先切断电源再灭火。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第三节", title="安全技术措施——一通三防",
            content="\n".join(ch8s3_lines), source="template",
        ))

        # 第八章-第四节 爆破
        ch8s4_lines = [
            "第四节  爆破",
            "",
            "一、一般规定",
            "  （一）爆破作业必须执行'一炮三检'和'三人连锁放炮'制度。",
            "  （二）爆破前瓦斯浓度≥0.5%时严禁爆破。",
            "  （三）爆破母线必须使用专用放炮线，长度不短于规定值。",
            f"  （四）掘进方式：{params_dict.get('dig_method', '—')}——"
            + ("综掘工作面一般不进行爆破作业，特殊情况除外。" if params_dict.get('dig_method') == '综掘' else "炮掘工作面必须严格执行爆破安全管理规定。"),
            "",
            "二、爆破器材管理",
            "  （一）炸药、雷管必须分开存放、分开运输。",
            "  （二）使用后的爆破器材必须当班退回，严禁私藏。",
            "  （三）爆破器材的领用、使用、退还必须按规定登记。",
            "",
            "三、爆破操作规程",
            "  （一）装药前必须检查炮眼深度、角度和间距。",
            "  （二）封泥长度不得小于200mm，水炮泥外剩余部分用粘土炮泥封实。",
            "  （三）爆破后必须通风排烟≥15分钟后方可进入工作面检查。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第四节", title="安全技术措施——爆破",
            content="\n".join(ch8s4_lines), source="template",
        ))

        # 第八章-第五节 机电
        ch8s5_lines = [
            "第五节  机电",
            "",
            "一、一般规定",
            "  （一）所有井下电气设备必须取得煤矿矿用产品安全标志。",
            f"  （二）掘进设备：{params_dict.get('dig_equipment', '—')}",
            "  （三）操作人员必须经专业培训持证上岗。",
            "  （四）设备检修必须切断电源并悬挂停电牌。",
            "",
            "二、掘进机操作规程",
            "  （一）开机前检查：截割头、液压系统、电缆、水路等。",
            "  （二）截割顺序：先顶后底，先中间后两帮。",
            "  （三）截割时严禁人员在截割臂回转范围内。",
            "  （四）掘进机运行中严禁检修和调整截割头。",
            "",
            "三、刮板输送机管理",
            "  （一）开机前必须检查链条紧度、刮板完好情况。",
            "  （二）运行中严禁人员跨越或在机上行走。",
            "  （三）机头、机尾必须设置急停装置。",
            "",
            "四、皮带输送机管理",
            "  （一）必须装设打滑保护、堆煤保护、防跑偏保护、温度保护、烟雾保护。",
            "  （二）皮带接头强度不低于皮带本身强度的85%。",
            "",
            "五、电气设备管理",
            "  （一）各部位绝缘电阻值应不低于0.5MΩ。",
            "  （二）接地极接地电阻不超过2Ω。",
            "  （三）检漏继电器动作电阻≤11kΩ，漏电闭锁动作可靠。",
            "  （四）供电系统必须实现三专（专用变压器、专用开关、专用电缆）。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第五节", title="安全技术措施——机电",
            content="\n".join(ch8s5_lines), source="template",
        ))

        # 第八章-第六节 运输
        ch8s6_lines = [
            "第六节  运输",
            "",
            f"一、运输方式：{params_dict.get('transport_method', '—')}",
            "",
            "二、一般规定",
            "  （一）各种运输设备运行前必须发出声光警报信号。",
            "  （二）运输设备的安全保护装置必须齐全、动作可靠。",
            "  （三）人员通过运输设备时必须停机。",
            "",
            "三、主运输设备管理",
            "  （一）带式输送机运行速度不得超过额定速度。",
            "  （二）带式输送机巷道每隔100m设置一个急停按钮。",
            "",
            "四、辅助运输管理",
            "  （一）轨道运输限速和信号联络制度。",
            "  （二）倾斜巷道运输防跑车装置：挡车栏、阻车器。",
            "  （三）无轨胶轮车限速和避让规定。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第六节", title="安全技术措施——运输",
            content="\n".join(ch8s6_lines), source="template",
        ))

        # 第八章-第七节 监控与通讯
        ch8s7_lines = [
            "第七节  监控与通讯",
            "",
            "一、安全监控系统",
            "  （一）传感器布置要求：",
            f"    1. 甲烷传感器T1：安设在距掘进工作面迎头≤5m处",
            f"       报警浓度≥0.5%CH₄，断电浓度≥{'0.5' if gas_level == '突出' else '1.0'}%CH₄",
            f"       复电浓度<{'0.5' if gas_level == '突出' else '1.0'}%CH₄",
            f"    2. 甲烷传感器T2：安设在回风流中距回风口10~15m处",
            "       报警浓度≥0.5%CH₄",
            "    3. 一氧化碳传感器：安设在回风流中",
            "       报警浓度≥24ppm",
            "    4. 风速传感器：安设在回风流中",
            "    5. 温度传感器：安设在掘进工作面",
            "",
            f"  （二）{'高瓦斯矿井使用量程上限不低于10%的甲烷传感器。' if gas_level in ('高瓦斯', '突出') else ''}",
            f"  {'（三）突出煤层使用量程上限不低于40%的甲烷传感器。' if gas_level == '突出' else ''}",
            "",
            "二、人员定位系统",
            "  井下作业人员必须佩戴人员位置识别卡。",
            "",
            "三、通讯系统",
            "  （一）掘进工作面必须安设有线调度电话。",
            "  （二）电话安设位置：距掘进面30~50m。",
            "",
            "四、视频监控",
            "  （一）掘进工作面迎头安设1台防爆摄像仪。",
            "  （二）视频监控存储≥30天。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第七节", title="安全技术措施——监控与通讯",
            content="\n".join(ch8s7_lines), source="template",
        ))

        # 第八章-第八节 防治水
        ch8s8_lines = [
            "第八节  防治水",
            "",
            f"一、水文地质类型：{params_dict.get('hydro_type', '—')}",
            "",
            "二、探放水管理",
            "  （一）坚持'有疑必探、先探后掘'的原则。",
            "  （二）探放水钻孔参数：孔径75mm，超前距离不少于规定值。",
            "  （三）每次探水前必须制定探放水设计。",
            "",
            "三、防治水措施",
            "  （一）工作面必须配备排水设备（水泵、管路），排水能力大于最大涌水量。",
            "  （二）掘进巷道低洼处必须设临时水仓。",
            "  （三）井下突水征兆：煤壁出汗、顶板淋水加大、水色发浑、空气变冷。",
            "",
            "四、应急处置",
            "  （一）发现突水征兆立即停止作业、报告调度室、撤出人员。",
            "  （二）撤退路线避开低洼处，向高处行进。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第八节", title="安全技术措施——防治水",
            content="\n".join(ch8s8_lines), source="template",
        ))

        # 第八章-第九节 其他
        ch8s9_lines = [
            "第九节  其他",
            "",
            "一、安全生产标准化",
            "  （一）掘进工作面必须达到安全生产标准化要求。",
            "  （二）巷道净断面不得小于设计断面的93%。",
            "  （三）锚杆、锚索支护质量必须逐根检查。",
            "",
            "二、文明生产",
            "  （一）巷道内物料码放整齐，不影响行人和运输。",
            "  （二）管线悬挂整齐，标识清晰。",
            "  （三）工作面保持清洁，无积水、无杂物。",
            "",
            "三、其他安全技术措施",
            "  （一）过地质构造段、应力集中区等特殊地段的安全措施。",
            "  （二）设备安装、拆除期间的安全措施。",
            "  （三）季节性安全防范措施（雨季防治水、冬季防火防冻等）。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第八章第九节", title="安全技术措施——其他",
            content="\n".join(ch8s9_lines), source="template",
        ))

        # ================================================================
        #  第九章  安全风险管控及应急避险 — 拆分为 3 个子章节
        # ================================================================

        ch9s1_lines = [
            "第一部分  安全风险管控",
            "",
            "一、安全风险辨识评估",
            "  按照安全风险专项辨识评估资料，重点描述工作面作业环境、",
            "  工程技术、设备设施、现场操作等方面存在的安全风险及相应管控措施。",
            f"  瓦斯等级：{gas_level}",
            f"  自燃倾向性：{params_dict.get('spontaneous_combustion', '—')}",
            f"  水文地质类型：{params_dict.get('hydro_type', '—')}",
            f"  围岩级别：{params_dict.get('rock_class', '—')}",
            "",
            "二、重大安全风险管控",
            "  （一）顶板风险：围岩破碎段、交岔点、开口段",
            "  （二）瓦斯风险：过构造段、揭煤期间",
            "  （三）水害风险：充水含水层、老空水",
            "  （四）机电运输风险：皮带运输、掘进机截割",
            "",
            "三、安全风险管控措施表",
            "  每条须包含：风险描述→风险等级(R=L×S)→管控措施→责任岗位→检查频次",
        ]
        chapters.append(ChapterContent(
            chapter_no="第九章-风险管控", title="安全风险管控",
            content="\n".join(ch9s1_lines), source="template",
        ))

        ch9s2_lines = [
            "第二部分  紧急避险设施",
            "",
            "一、自救器配备",
            "  掘进工作面所有作业人员必须随身携带自救器（隔绝式化学氧45min型）。",
            "  佩戴方法培训和演练每季度≥1次。",
            "",
            "二、压风自救系统",
            "  压风自救装置安设间距不超过200m。",
            "  压风管路规格：主管Φ150mm，支管Φ25mm。",
            "  每组压风自救装置供≥8人使用。",
            "",
            "三、供水施救系统",
            "  供水施救管路与压风管路同侧敷设。",
            "  每隔50m设置三通阀门。",
            "",
            "四、避难硐室/避难所",
            "  距工作面500m以内设置临时避难硐室或移动式救生舱。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第九章-紧急避险", title="紧急避险设施",
            content="\n".join(ch9s2_lines), source="template",
        ))

        ch9s3_lines = [
            "第三部分  灾害应急处置",
            "",
            "一、应急响应程序",
            "  明确授予带班人员、班组长、瓦检工、调度人员遇险处置权和紧急避险权。",
            "",
            "二、各类灾害应急处置措施",
            "  （一）顶板事故应急处置：加固支护→撤人→汇报→救援",
            "  （二）瓦斯超限/爆炸处置：切断电源→撤人→反风→救援",
            "  （三）煤尘爆炸处置：利用隔爆设施→撤离路线→急救",
            "  （四）火灾处置：切断电源→直接灭火/封闭灭火→撤人",
            "  （五）水灾处置：向高处撤离→报告调度→启动排水预案",
            "",
            "三、避灾路线（分灾种标识）",
            "  水灾撤退路线：掘进面→进风巷→采区上山→地面",
            "  火灾/瓦斯撤退路线：掘进面→回风巷→总回风→地面",
            "  （详见附图：避灾路线图）",
            "",
            "四、救护演练",
            "  每季度组织≥1次应急救援演练。",
            "  演练内容：撤人路线、自救器使用、避难硐室启用。",
        ]
        chapters.append(ChapterContent(
            chapter_no="第九章-应急处置", title="灾害应急处置",
            content="\n".join(ch9s3_lines), source="template",
        ))

        # ================================================================
        #  第十章  主要设备及材料清单
        # ================================================================
        if equipment_result:
            # 设备清单章节
            eq_lines = ["一、主要设备配备表", ""]
            eq_lines.append(f"  本工作面配备主要设备共 {equipment_result.total_equipment_count} 台（套），详见下表：")
            eq_lines.append("")
            eq_lines.append("  序号 | 设备类别 | 设备名称 | 型号规格 | 数量 | 功率(kW)")
            eq_lines.append("  --- | --- | --- | --- | --- | ---")
            for i, eq in enumerate(equipment_result.equipment_list, 1):
                power_str = f"{eq.power_kw}" if eq.power_kw else "—"
                eq_lines.append(
                    f"  {i} | {eq.category} | {eq.name} | {eq.model_spec or '—'} | {eq.quantity} | {power_str}"
                )
            eq_lines.append("")
            # 总装机功率
            total_power = sum(eq.power_kw or 0 for eq in equipment_result.equipment_list)
            if total_power > 0:
                eq_lines.append(f"  总装机功率：{total_power:.1f} kW")

            chapters.append(ChapterContent(
                chapter_no="第十章-设备", title="主要设备配备",
                content="\n".join(eq_lines), source="equipment_match",
            ))

            # 材料工程量章节
            mat_lines = ["二、主要材料工程量清单", ""]
            mat_lines.append(f"  本工作面共需 {equipment_result.total_material_types} 种主要支护材料，详见下表：")
            mat_lines.append("")
            mat_lines.append("  序号 | 材料类别 | 材料名称 | 规格型号 | 单位 | 单循环 | 月用量 | 工程总量")
            mat_lines.append("  --- | --- | --- | --- | --- | --- | --- | ---")
            for i, mat in enumerate(equipment_result.material_bom, 1):
                mat_lines.append(
                    f"  {i} | {mat.category} | {mat.name} | {mat.model_spec or '—'} | "
                    f"{mat.unit} | {mat.qty_per_cycle or '—'} | {mat.qty_per_month or '—'} | {mat.qty_total or '—'}"
                )
            mat_lines.append("")
            if equipment_result.material_bom:
                mat_lines.append(f"  计算依据：{equipment_result.material_bom[0].calc_basis or '—'}")

            chapters.append(ChapterContent(
                chapter_no="第十章-材料", title="主要材料工程量清单",
                content="\n".join(mat_lines), source="equipment_match",
            ))
        else:
            # 降级：无设备材料匹配结果时仍生成占位章节
            chapters.append(ChapterContent(
                chapter_no="第十章", title="主要设备及材料清单",
                content="（设备材料配置数据待补充，请通过设备材料管理页面完成配置后重新生成。）",
                source="template",
            ))

        # ================================================================
        #  附录：编制依据与规则命中
        # ================================================================
        if match_result and match_result.matched_rules:
            rule_lines = []
            for mr in match_result.matched_rules:
                rule_lines.append(f"• {mr.rule_name}（{mr.category}，优先级 {mr.priority}）")
                for a in mr.actions:
                    rule_lines.append(f"  → 关联章节：{a.target_chapter}")

            chapters.append(ChapterContent(
                chapter_no="附录", title="编制依据与规则命中",
                content="\n".join(rule_lines), source="rule_match",
            ))

        return chapters


    async def _ai_polish_content(
        self,
        chapters: list[ChapterContent],
        params: dict,
        calc_result=None,
        vent_result=None,
        tenant_id: int = 0,
        project=None,
    ) -> list[ChapterContent]:
        """
        AI 深度扩写引擎 — 交付顶级版（范文 Few-shot + 计算推导 + 多轮扩写）

        核心策略:
          1. 范文 Few-shot：从 reference_chapters.json 加载对应章节作为风格示例
          2. 计算推导注入：支护/通风计算引擎的完整公式推导文本注入对应章节
          3. RAG 全量检索：top_k=20/25, 阈值 0.3, 每条截取 3000 字
          4. 多轮扩写：第一轮不限 max_tokens，不足 6000 字自动第二轮
          5. 每节目标 8000+ 字符，超越范文水平
        """
        from app.core.config import settings
        from openai import AsyncOpenAI
        from app.services.embedding_service import EmbeddingService
        import json, os

        api_key = settings.OPENAI_API_KEY or settings.GEMINI_API_KEY
        base_url = settings.OPENAI_BASE_URL or None
        model = settings.AI_MODEL

        if not api_key:
            return chapters

        client_kwargs = {"api_key": api_key}
        if base_url:
            client_kwargs["base_url"] = base_url
        client = AsyncOpenAI(**client_kwargs)

        emb_svc = EmbeddingService(self.session)

        # ===== 构建真值锚点（不可编造的硬数据） =====
        face_name = getattr(project, 'face_name', '') if project else ''
        mine_name = getattr(project, 'mine_name', '') if project else ''
        anchor_parts = []
        if face_name:
            anchor_parts.append(f"巷道编号: {face_name}")
        if mine_name:
            anchor_parts.append(f"矿井名称: {mine_name}")
        _anchor_keys = {
            'section_width': '断面宽度(m)', 'section_height': '断面高度(m)',
            'excavation_length': '掘进长度(m)', 'gas_level': '瓦斯等级',
            'rock_class': '围岩类别', 'section_form': '断面形式',
            'coal_seam': '煤层名称', 'coal_thickness': '煤层厚度(m)',
            'seam_dip': '煤层倾角(°)', 'dig_method': '掘进方式',
            'dig_equipment': '掘进设备', 'transport_method': '运输方式',
        }
        for k, label in _anchor_keys.items():
            v = params.get(k)
            if v is not None and str(v).strip():
                anchor_parts.append(f"{label}: {v}")
        truth_anchor = "\n".join(anchor_parts)
        self._truth_anchor = truth_anchor
        if truth_anchor:
            print(f"⚓ 真值锚点构建完成: {len(anchor_parts)} 项参数")

        # ===== 加载范文片段（Few-shot 示例） =====
        ref_chapters: dict[str, str] = {}
        ref_path = os.path.join(
            os.path.dirname(__file__), '..', 'data', 'reference_chapters.json'
        )
        try:
            with open(ref_path, 'r', encoding='utf-8') as f:
                ref_chapters = json.load(f)
            print(f"📖 已加载范文片段: {len(ref_chapters)} 节")
        except Exception as e:
            print(f"⚠️ 范文片段加载失败: {e}")

        # ===== 生成计算推导文本 =====
        support_narrative = ""
        vent_narrative = ""
        if calc_result:
            try:
                from app.services.calc_engine import SupportCalcEngine
                from app.schemas.calc import SupportCalcInput
                # 重建输入参数
                calc_inp = SupportCalcInput(
                    rock_class=params.get('rock_class', 'III'),
                    section_width=float(params.get('section_width', 5.0)),
                    section_height=float(params.get('section_height', 3.5)),
                    section_form=params.get('section_form', '矩形'),
                    rock_density=float(params.get('rock_density', 2.5)),
                    bolt_length=float(params.get('bolt_length', 2.0)),
                    bolt_diameter=float(params.get('bolt_diameter', 20)),
                    cable_strength=float(params.get('cable_strength', 260)),
                )
                support_narrative = SupportCalcEngine.narrative(calc_inp, calc_result)
                print(f"📐 支护计算推导: {len(support_narrative)} 字符")
            except Exception as e:
                print(f"⚠️ 支护推导生成失败: {e}")

        if vent_result:
            try:
                from app.services.vent_engine import VentCalcEngine
                from app.schemas.vent import VentCalcInput
                calc_inp = VentCalcInput(
                    gas_level=params.get('gas_level', '低瓦斯'),
                    gas_emission=float(params.get('gas_emission', 1.0)),
                    max_workers=int(params.get('max_workers', 20)),
                    explosive_per_cycle=float(params.get('explosive_per_cycle', 0)),
                    section_area=float(params.get('section_width', 5.0)) * float(params.get('section_height', 3.5)),
                    excavation_length=float(params.get('excavation_length', 1000)),
                )
                vent_narrative = VentCalcEngine.narrative(calc_inp, vent_result)
                print(f"🌬️ 通风计算推导: {len(vent_narrative)} 字符")
            except Exception as e:
                print(f"⚠️ 通风推导生成失败: {e}")

        ROUND1_TIMEOUT = 240   # 加长超时，确保大章节不被截断
        ROUND2_TIMEOUT = 150
        MIN_CHARS_FOR_ROUND2 = 6000

        SYSTEM_PROMPT = """你是中国煤矿安全领域最顶尖的规程编制专家，曾主编多部国家级煤矿安全技术标准。
你拥有 30 年以上井下一线安全管理和规程编审经验，熟悉华阳集团所有技术管理制度。

你必须严格遵循以下法规和企业标准：
1. 《煤矿安全规程》（2022版）— 国家矿山安全监察局
2. 《煤矿安全生产标准化管理体系基本要求及评分方法》矿安[2024]109号
3. 《煤矿安全生产标准化管理体系考核定级办法》矿安[2024]109号
4. 《煤矿作业规程编制指南》煤炭工业出版社（2024年版）
5. 《煤炭生产技术管理规定》华阳发〔2025〕242号
6. 《华阳集团安全管理工作制度》华阳安字〔2024〕431号
7. 《"一通三防"工作管理制度》华阳发〔2024〕119号
8. 《华阳集团机电管理制度》华阳股份发〔2024〕116号
9. 《华阳集团地测防治水工作管理制度》华阳发〔2024〕345号
10. 《工业视频监控系统管理规定（试行）》华阳发〔2024〕94号
11. 《山西省煤矿顶板安全管理规定》晋应急发〔2019〕299号
12. 《山西省煤矿防治水"三专两探一撤"规定》晋应急发〔2019〕270号

== 输出铁律 ==
- 严禁使用任何模糊表述："按规定执行""根据实际情况""按要求操作"——必须替换为具体数值和操作步骤
- 每条措施必须有【具体数值标准】和【责任岗位】
- 必须引用法规全称及条款编号
- 按华阳集团规程编制大纲层级：章→节→条→款→项
- 技术参数标注单位和取值依据
- 操作步骤有明确的先后顺序
- 安全措施按"一般规定→具体操作→质量标准→注意事项→应急处置"展开
- 计算校核必须有完整公式推导和数值代入过程
- 质量标准以量化指标形式给出（不得定性描述）

== 数据引用铁律（违反此规则视为严重错误） ==
- 下文「真值锚点」中列出的工程参数是唯一数据源，正文必须原样引用，严禁修改数值或编造替代值
- 巷道编号必须全文统一使用「真值锚点」中给出的名称，严禁使用任何其他巷道编号
- 如果某个技术参数在「真值锚点」和「参数基线」中都没有提供，必须使用占位符「【待补充】」，严禁自行编造坐标、标高、断面尺寸等硬数据
- 同一参数在不同章节出现时，数值必须完全一致

== 公式格式铁律 ==
- 公式输出必须使用纯文本格式（如 Q = K × S = 1.2 × 12.0 = 14.4 m³/min）
- 严禁使用 LaTeX 数学语法（如 $Q=KS$ 或 \\frac{} 等），Word 文档无法渲染 LaTeX
- 上下标用文字说明（如 V风速 而不是 V_{风速}）"""

        async def _polish_one(ch: ChapterContent) -> None:
            """单章多轮深度扩写"""

            if ch.chapter_no == "附录":
                return

            # ===== 全量 RAG 检索 =====
            rag_context_parts = []
            import re

            # 多角度查询
            queries = [
                ch.title,
                f"掘进工作面 {ch.title}",
                f"煤矿 {ch.title} 安全技术措施",
            ]
            section_titles = re.findall(r'第[一二三四五六七八九十]+节\s*(.+)', ch.content)
            for st in section_titles[:5]:
                queries.append(st.strip())
            # 从内容中提取关键词条
            keywords = re.findall(r'[一二三四五六七八九十]+、\s*(.+)', ch.content)
            for kw in keywords[:5]:
                queries.append(kw.strip()[:20])

            # 标准库 — 全量检索
            try:
                seen = set()
                for q in queries:
                    results = await emb_svc.search_similar(
                        query=q, tenant_id=tenant_id, top_k=20, threshold=0.3
                    )
                    if results:
                        for r in results:
                            key = r['content'][:80]
                            if key not in seen:
                                seen.add(key)
                                rag_context_parts.append(
                                    f"[标准-{r['doc_title']}] {r['clause_no']}:\n{r['content'][:3000]}"
                                )
            except Exception:
                pass

            # 知识库 — 全量检索
            try:
                seen_s = set()
                for q in queries:
                    results = await emb_svc.search_snippets(
                        query=q, tenant_id=tenant_id, top_k=25, threshold=0.3
                    )
                    if results:
                        for r in results:
                            key = r['content'][:80]
                            if key not in seen_s:
                                seen_s.add(key)
                                rag_context_parts.append(
                                    f"[客户规程-{r.get('chapter_name', '')}]:\n{r['content'][:3000]}"
                                )
            except Exception:
                pass

            # ===== RAG 相关性过滤 =====
            # 按章节标题关键词相关性排序，高相关排前面
            if rag_context_parts:
                ch_keywords = set(ch.title) | set(ch.chapter_no)
                def _relevance(text):
                    return sum(1 for kw in ch_keywords if kw in text[:200])
                rag_context_parts.sort(key=_relevance, reverse=True)
                # 保留前 30 条高相关结果
                rag_context_parts = rag_context_parts[:30]

            rag_context = "\n\n---\n\n".join(rag_context_parts)

            # ===== 第一轮：完整生成 =====
            # ===== 组装动态 Prompt 变量 =====
            ch_baseline = getattr(ch, '_baseline', '')
            baseline_text = f"== 前序核心章节参数基线（引用数据必须与此一致） ==\n{ch_baseline[:6000]}\n\n" if ch_baseline else ""

            # 真值锚点注入（拼在 baseline 前面，确保每轮都能看到）
            anchor_text = getattr(self, '_truth_anchor', '')
            anchor_section = f"== 真值锚点（以下工程参数为唯一数据源，严禁修改或编造） ==\n{anchor_text}\n\n" if anchor_text else ""
            baseline_text = anchor_section + baseline_text

            rag_context_text = f"== 标准库和客户已有规程参考资料（务必深度融合） ==\n{rag_context}\n\n" if rag_context else ""
            
            fewshot_raw = self._find_reference_section(ch, ref_chapters)
            few_shot_text = f"== 范文示例（对标此风格和专业深度，输出必须超越此水平） ==\n{fewshot_raw}\n\n" if fewshot_raw else ""
            
            ch_lower = f"{ch.chapter_no} {ch.title}".lower()
            calc_parts = []
            if support_narrative and any(k in ch_lower for k in ['支护', '顶板控制', '顶板管理', '第三章', '第四章']):
                calc_parts.append(f"== 支护计算推导过程（必须完整引用到正文中） ==\n{support_narrative}\n")
            if vent_narrative and any(k in ch_lower for k in ['通风', '一通三防', '生产系统', '第五章']):
                calc_parts.append(f"== 通风计算推导过程（必须完整引用到正文中） ==\n{vent_narrative}\n")
            calculation_text = "\n".join(calc_parts) + "\n" if calc_parts else ""

            # ===== 动态加载第一轮生成 Prompt =====
            prompt_r1 = prompt_manager.format_prompt(
                category="doc_generation",
                version_key="v2_few_shot",
                chapter_no=ch.chapter_no,
                title=ch.title,
                params_text=self._format_params_for_prompt(params),
                outline=ch.content,
                baseline_text=baseline_text,
                rag_context_text=rag_context_text,
                calculation_text=calculation_text,
                few_shot_text=few_shot_text
            )

            try:
                resp = await asyncio.wait_for(
                    client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": SYSTEM_PROMPT},
                            {"role": "user", "content": prompt_r1},
                        ],
                        # 不设 max_tokens — 让模型自由输出完整内容
                    ),
                    timeout=ROUND1_TIMEOUT,
                )
                r1_content = resp.choices[0].message.content or ""
                r1_len = len(r1_content)
                print(f"📝 第一轮完成: {ch.chapter_no} {ch.title} → {r1_len} 字符")

                if r1_len > len(ch.content):
                    ch.content = r1_content
                    ch.source = "ai_polished"

                    # ===== 第二轮：补充扩展（如不足 6000 字） =====
                    if r1_len < MIN_CHARS_FOR_ROUND2:
                        prompt_r2 = (
                            f"以下是你刚才编写的【{ch.chapter_no} {ch.title}】章节内容，"
                            f"但内容仍不够详尽（当前仅 {r1_len} 字符）。\n\n"
                            f"请基于以下已有内容，进行【深度补充扩展】：\n"
                            f"1. 对每一个条款增加更多的子款和具体操作步骤\n"
                            f"2. 补充遗漏的安全技术要求（如设备操作规程、特殊情况处理等）\n"
                            f"3. 增加质量验收标准和检查频次\n"
                            f"4. 补充应急处置措施\n\n"
                            f"已有内容：\n{r1_content}\n\n"
                            f"请输出【补充扩展内容】，直接续接在已有内容之后。"
                            f"输出不少于 3000 字。不要重复已有内容。"
                        )
                        try:
                            resp2 = await asyncio.wait_for(
                                client.chat.completions.create(
                                    model=model,
                                    messages=[
                                        {"role": "system", "content": SYSTEM_PROMPT},
                                        {"role": "user", "content": prompt_r2},
                                    ],
                                ),
                                timeout=ROUND2_TIMEOUT,
                            )
                            r2_content = resp2.choices[0].message.content or ""
                            if r2_content and len(r2_content) > 500:
                                ch.content = r1_content + "\n\n" + r2_content
                                print(f"📝 第二轮补充: {ch.chapter_no} +{len(r2_content)} 字符 → 总计 {len(ch.content)} 字符")
                        except asyncio.TimeoutError:
                            print(f"⏱️ 第二轮超时: {ch.title}（保留第一轮结果）")
                        except Exception as e2:
                            print(f"⚠️ 第二轮失败: {ch.title}: {e2}")

                    # ===== 第三轮：Critic 自我评估 =====
                    CRITIC_TIMEOUT = 90
                    current_content = ch.content
                    import re as _re

                    critic_prompt = prompt_manager.format_prompt(
                        category="critic_evaluation",
                        chapter_no=ch.chapter_no,
                        title=ch.title,
                        content=current_content[:8000]
                    )

                    try:
                        critic_resp = await asyncio.wait_for(
                            client.chat.completions.create(
                                model=model,
                                messages=[
                                    {"role": "system", "content": "你是中国煤矿安全规程的质量审核专家，审核标准极其严格。"},
                                    {"role": "user", "content": critic_prompt},
                                ],
                            ),
                            timeout=CRITIC_TIMEOUT,
                        )
                        critic_output = critic_resp.choices[0].message.content or ""

                        # 解析 Critic 结果
                        has_issues = (
                            "PASS" not in critic_output
                            and "ISSUES_COUNT: 0" not in critic_output
                            and "FIXES:" in critic_output
                        )

                        if has_issues:
                            # 提取评分
                            score_m = _re.search(r'SCORE:\s*(\d+)', critic_output)
                            score = int(score_m.group(1)) if score_m else 5
                            print(f"🔍 Critic 评估: {ch.chapter_no} → {score}/10, 存在问题需修正")

                            # ===== 第四轮：定向修正 =====
                            FIX_TIMEOUT = 150
                            fix_prompt = (
                                f"以下是【{ch.chapter_no} {ch.title}】章节的审核报告和原文。\n"
                                f"请根据审核报告中指出的问题，对原文进行【精准修正】。\n\n"
                                f"修正要求：\n"
                                f"1. 将所有模糊表述替换为具体的数值标准和操作步骤\n"
                                f"2. 补充缺失的法规引用（引用全称+条款号）\n"
                                f"3. 补充缺失的责任岗位\n"
                                f"4. 补全缺失的结构块（质量标准/应急处置等）\n"
                                f"5. 保留原文中已经合格的内容，只修改有问题的部分\n\n"
                                f"== 审核报告 ==\n{critic_output}\n\n"
                                f"== 原文 ==\n{current_content}\n\n"
                                f"直接输出修正后的完整章节正文。不要输出任何注释或说明。"
                            )

                            try:
                                fix_resp = await asyncio.wait_for(
                                    client.chat.completions.create(
                                        model=model,
                                        messages=[
                                            {"role": "system", "content": SYSTEM_PROMPT},
                                            {"role": "user", "content": fix_prompt},
                                        ],
                                    ),
                                    timeout=FIX_TIMEOUT,
                                )
                                fixed = fix_resp.choices[0].message.content or ""
                                if fixed and len(fixed) > len(current_content) * 0.8:
                                    ch.content = fixed
                                    ch.source = "ai_self_iterated"
                                    print(f"🔧 定向修正完成: {ch.chapter_no} {score}/10 → {len(fixed)} 字符")
                                else:
                                    print(f"⚠️ 修正输出过短, 保留原文: {ch.title}")
                            except asyncio.TimeoutError:
                                print(f"⏱️ 定向修正超时: {ch.title}（保留 Critic 前内容）")
                            except Exception as ef:
                                print(f"⚠️ 定向修正失败: {ch.title}: {ef}")
                        else:
                            score_m = _re.search(r'SCORE:\s*(\d+)', critic_output) if 'SCORE' in critic_output else None
                            score = int(score_m.group(1)) if score_m else 8
                            print(f"✨ Critic 通过: {ch.chapter_no} → {score}/10, 无需修正")

                    except asyncio.TimeoutError:
                        print(f"⏱️ Critic 评估超时: {ch.title}（跳过自评）")
                    except Exception as ec:
                        print(f"⚠️ Critic 评估失败: {ch.title}: {ec}")

                    print(f"✅ AI 扩写完成: {ch.chapter_no} {ch.title} → {len(ch.content)} 字符")
                else:
                    print(f"⚠️ AI 输出过短, 保留模板: {ch.title}")

            except asyncio.TimeoutError:
                print(f"⏱️ AI 扩写超时({ROUND1_TIMEOUT}s), 降级使用模板: {ch.title}")
            except Exception as e:
                print(f"⚠️ AI 扩写失败: {ch.title}: {e}")

            # ===== AI 输出后处理清洗（防御性修正） =====
            if ch.source in ('ai_polished', 'ai_self_iterated', 'ai_critic_fixed'):
                ch.content = self._postprocess_ai_output(ch.content, face_name)

        # ===== 半串行三阶段架构 =====

        # 核心章节标识（概述、地质、支护 — 包含最基础工程数据）
        CORE_PREFIXES = ('第一章', '第二章', '第三章')
        core_chapters = [ch for ch in chapters if ch.chapter_no.startswith(CORE_PREFIXES)]
        other_chapters = [ch for ch in chapters if not ch.chapter_no.startswith(CORE_PREFIXES)]

        # --- Phase 1: 核心章节优先并发生成（3 章并发远快于串行） ---
        print(f"📌 Phase1: 优先生成 {len(core_chapters)} 个核心章节...")
        await asyncio.gather(*[_polish_one(ch) for ch in core_chapters])

        # --- 提取参数基线（从核心章节提取关键数据） ---
        baseline_parts = []
        for ch in core_chapters:
            if ch.source in ('ai_polished', 'ai_self_iterated') and len(ch.content) > 500:
                # 提取前 1500 字作为基线摘要
                baseline_parts.append(
                    f"[{ch.chapter_no} {ch.title} 摘要]:\n{ch.content[:1500]}"
                )
        param_baseline = "\n\n".join(baseline_parts)
        if param_baseline:
            print(f"📋 参数基线提取完成: {len(param_baseline)} 字符, 覆盖 {len(baseline_parts)} 章节")

        # 将基线注入后续章节的 prompt（通过闭包变量共享）
        # _polish_one 中需要用到 param_baseline，通过 nonlocal 不方便
        # 所以在每个 other_chapter 的 content 开头临时注入标记
        BASELINE_MARKER = "\n\n== 前序章节参数基线（必须保持一致） ==\n"
        for ch in other_chapters:
            if param_baseline and ch.chapter_no != "附录":
                ch._baseline = param_baseline  # 动态属性存储基线

        # --- Phase 2: 其余章节并发生成（带基线上下文） ---
        print(f"📌 Phase2: 并发生成 {len(other_chapters)} 个后续章节...")
        await asyncio.gather(*[_polish_one(ch) for ch in other_chapters])

        # --- Phase 3: 跨章节一致性扫描 ---
        print("📌 Phase3: 跨章节一致性扫描...")
        try:
            # 收集全文关键数据摘要
            all_summaries = []
            for ch in chapters:
                if ch.chapter_no != "附录" and len(ch.content) > 200:
                    all_summaries.append(
                        f"[{ch.chapter_no} {ch.title}]: {ch.content[:800]}"
                    )

            consistency_prompt = prompt_manager.format_prompt(
                category="consistency_scan",
                summaries='\n\n'.join(all_summaries[:15])
            )

            consistency_resp = await asyncio.wait_for(
                client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "你是煤矿安全规程数据一致性审查专家。"},
                        {"role": "user", "content": consistency_prompt},
                    ],
                ),
                timeout=90,
            )
            consistency_result = consistency_resp.choices[0].message.content or ""

            if "CONSISTENT" in consistency_result and "INCONSISTENCY" not in consistency_result:
                print("✅ 跨章节一致性检查通过：无矛盾")
            else:
                # 发现矛盾，输出日志（当前版本记录报告，后续可自动修正）
                inconsistencies = [
                    line for line in consistency_result.split('\n')
                    if 'INCONSISTENCY' in line
                ]
                print(f"⚠️ 发现 {len(inconsistencies)} 处跨章节矛盾:")
                for inc in inconsistencies[:5]:
                    print(f"   {inc}")

                # 自动修正：找到矛盾章节，用正确数据修正
                if inconsistencies:
                    fix_prompt = (
                        '以下是规程中发现的数据矛盾，以及需要修正的章节原文。\n'
                        '请根据矛盾报告，修正原文中的错误数据，使其与正确数据一致。\n'
                        '只修改有错误的数据，保留其他内容不变。\n\n'
                        f'== 矛盾报告 ==\n{consistency_result}\n\n'
                    )
                    # 对受影响的非核心章节做轻量修正
                    for ch in other_chapters:
                        affected = any(
                            ch.chapter_no in inc or ch.title[:4] in inc
                            for inc in inconsistencies
                        )
                        if affected and len(ch.content) > 500:
                            try:
                                fix_resp = await asyncio.wait_for(
                                    client.chat.completions.create(
                                        model=model,
                                        messages=[
                                            {"role": "system", "content": SYSTEM_PROMPT},
                                            {"role": "user", "content": (
                                                fix_prompt
                                                + f'== 待修正章节: {ch.chapter_no} {ch.title} ==\n'
                                                + ch.content + '\n\n'
                                                + '直接输出修正后的完整章节正文。'
                                            )},
                                        ],
                                    ),
                                    timeout=150,
                                )
                                fixed = fix_resp.choices[0].message.content or ""
                                if fixed and len(fixed) > len(ch.content) * 0.8:
                                    ch.content = fixed
                                    ch.source = "ai_consistency_fixed"
                                    print(f"🔄 一致性修正: {ch.chapter_no} → {len(fixed)} 字符")
                            except Exception:
                                pass

        except asyncio.TimeoutError:
            print("⏱️ 一致性扫描超时（跳过）")
        except Exception as ec:
            print(f"⚠️ 一致性扫描失败: {ec}")

        return chapters

    @staticmethod
    def _postprocess_ai_output(content: str, face_name: str) -> str:
        """AI 输出后处理清洗 — LaTeX清洗 + 巷道编号强制替换 + 行去重"""
        import re as _re

        # 1. LaTeX 清洗：$...$ → 内部文本
        def _clean_latex(match):
            inner = match.group(1)
            inner = _re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', inner)
            inner = _re.sub(r'\\times', '×', inner)
            inner = _re.sub(r'\\div', '÷', inner)
            inner = _re.sub(r'\\leq', '≤', inner)
            inner = _re.sub(r'\\geq', '≥', inner)
            inner = _re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', inner)
            inner = _re.sub(r'\\[a-zA-Z]+', '', inner)
            inner = _re.sub(r'[{}]', '', inner)
            inner = inner.replace('^', '').replace('_', '')
            return inner.strip()
        content = _re.sub(r'\$([^$]+)\$', _clean_latex, content)

        # 2. 巷道编号强制替换
        if face_name:
            face_base = _re.search(r'(\d+)', face_name)
            if face_base:
                correct_num = face_base.group(1)
                def _fix_face_ref(m):
                    num, suffix = m.group(1), m.group(2)
                    if num != correct_num and suffix in face_name:
                        return face_name
                    return m.group(0)
                content = _re.sub(
                    r'(\d{4,5})(进风巷|回风巷|运输巷|切眼)',
                    _fix_face_ref, content
                )

        # 3. 章标题去重
        lines = content.split('\n')
        deduped = [lines[0]] if lines else []
        for i in range(1, len(lines)):
            if lines[i].strip() and lines[i].strip() == lines[i-1].strip():
                continue
            deduped.append(lines[i])
        return '\n'.join(deduped)

    @staticmethod
    def _find_reference_section(ch, ref_chapters: dict) -> str:
        """
        根据章节号和标题模糊匹配范文对应片段

        匹配策略：
          1. 精确匹配：第八章第三节 → "第八章第三节 一通三防"
          2. 关键词匹配：标题含"顶板管理"→ 找含"顶板"的范文键
          3. 章节号匹配：第五章 → 找"第五章"开头的所有片段合并
        """
        if not ref_chapters:
            return ""

        ch_id = ch.chapter_no  # 如 "第八章第三节"
        ch_title = ch.title    # 如 "安全技术措施——一通三防"

        # 策略 1: 精确匹配键前缀
        for key, val in ref_chapters.items():
            # "第八章第三节" 匹配 "第八章第三节 一通三防"
            if key.startswith(ch_id):
                return val[:5000]

        # 策略 2: 标题关键词匹配
        # 提取标题中的关键词
        import re
        keywords = re.split(r'[——\-\s]+', ch_title)
        keywords = [k for k in keywords if len(k) >= 2]
        for kw in keywords:
            for key, val in ref_chapters.items():
                if kw in key:
                    return val[:5000]

        # 策略 3: 章节号前缀匹配（合并多节）
        # "第五章" → 找所有 "第五章" 开头的
        ch_prefix = ch_id[:3] if len(ch_id) >= 3 else ch_id  # "第X章"
        parts = []
        for key, val in ref_chapters.items():
            if key.startswith(ch_prefix):
                parts.append(val[:2000])
        if parts:
            return "\n\n".join(parts[:3])[:5000]

        return ""

    @staticmethod
    def _format_params_for_prompt(params: dict) -> str:
        """将工程参数格式化为 Prompt 可读文本"""
        labels = {
            "rock_class": "围岩级别", "coal_thickness": "煤层厚度(m)",
            "coal_dip_angle": "煤层倾角(°)", "gas_level": "瓦斯等级",
            "hydro_type": "水文地质类型", "geo_structure": "地质构造",
            "spontaneous_combustion": "自燃倾向性", "roadway_type": "巷道类型",
            "excavation_type": "掘进类型", "section_form": "断面形式",
            "section_width": "断面宽度(m)", "section_height": "断面高度(m)",
            "excavation_length": "掘进长度(m)", "service_years": "服务年限(年)",
            "dig_method": "掘进方式", "dig_equipment": "掘进设备",
            "transport_method": "运输方式",
        }
        lines = []
        for k, v in params.items():
            if v is not None and k in labels:
                lines.append(f"  {labels[k]}: {v}")
        return "\n".join(lines) if lines else "  （参数未填写）"

    # ========== Word 渲染 ==========

    def _render_docx(self, project, chapters, calc_result, vent_result, equipment_result=None) -> str:
        """用 python-docx 生成 .docx 文件"""
        doc = Document()

        # 文档样式
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(12)

        # --- 封面 ---
        for _ in range(4):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run(f"{project.face_name}")
        run.font.size = Pt(22)
        run.font.bold = True

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2 = subtitle.add_run("掘进工作面作业规程")
        run2.font.size = Pt(18)

        doc.add_paragraph()

        meta_items = [
            f"矿井名称：{getattr(project, 'mine_name', '—')}",
            f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}",
            f"编制单位：生产技术科",
        ]
        for item in meta_items:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(item).font.size = Pt(14)

        doc.add_page_break()

        # --- 正文章节 ---
        import re as _docx_re

        # 子章节判定正则：第八章第X节、第九章-XXX 等拆分子节用 H2，主章节用 H1
        _SUB_CHAPTER_RE = _docx_re.compile(
            r'^第[一二三四五六七八九十\d]+章(?:第[一二三四五六七八九十\d]+节|[-—])'
        )
        # 内容中的节标题：仅匹配行首"第X节"后跟空格或汉字标题（避免误匹配正文段落）
        _SECTION_TITLE_RE = _docx_re.compile(
            r'^第[一二三四五六七八九十\d]+节[\s\u4e00-\u9fff]'
        )
        # 编号大项：一、二、……十、
        _NUMBERED_ITEM_RE = _docx_re.compile(r'^[一二三四五六七八九十]+[、．.]')
        # 条标题：第X条
        _CLAUSE_TITLE_RE = _docx_re.compile(r'^第[一二三四五六七八九十百\d]+条')
        # 数字条款 / 带括号编号：1. 2.（一）
        _SUB_ITEM_RE = _docx_re.compile(r'^(\d{1,2}[\.、]|（[一二三四五六七八九十]）)')

        for ch in chapters:
            # 章节标题层级
            _is_sub = bool(_SUB_CHAPTER_RE.match(ch.chapter_no))
            heading = doc.add_heading(f"{ch.chapter_no}  {ch.title}", level=2 if _is_sub else 1)

            # 预警标记
            if ch.has_warning:
                warn_p = doc.add_paragraph()
                warn_run = warn_p.add_run("⚠ 本章节存在合规预警，请重点审查")
                warn_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                warn_run.font.bold = True

            # ===== 设备/材料清单 → Word 表格渲染 =====
            if ch.source == "equipment_match" and equipment_result:
                if ch.chapter_no == "第十章-设备":
                    # 设备清单表格
                    doc.add_paragraph(
                        f"本工作面配备主要设备共 {equipment_result.total_equipment_count} 台（套），详见下表："
                    )
                    headers = ["序号", "设备类别", "设备名称", "型号规格", "数量", "功率(kW)"]
                    tbl = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
                    # 表头
                    for j, h in enumerate(headers):
                        cell = tbl.rows[0].cells[j]
                        cell.text = h
                        for run in cell.paragraphs[0].runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # 数据行
                    for i, eq in enumerate(equipment_result.equipment_list, 1):
                        row = tbl.add_row()
                        row.cells[0].text = str(i)
                        row.cells[1].text = eq.category
                        row.cells[2].text = eq.name
                        row.cells[3].text = eq.model_spec or "—"
                        row.cells[4].text = str(eq.quantity)
                        row.cells[5].text = f"{eq.power_kw}" if eq.power_kw else "—"
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(10)
                    # 总装机功率
                    total_power = sum(eq.power_kw or 0 for eq in equipment_result.equipment_list)
                    if total_power > 0:
                        p = doc.add_paragraph(f"总装机功率：{total_power:.1f} kW")
                        p.runs[0].font.bold = True
                    continue

                elif ch.chapter_no == "第十章-材料":
                    # 材料 BOM 表格
                    doc.add_paragraph(
                        f"本工作面共需 {equipment_result.total_material_types} 种主要支护材料，详见下表："
                    )
                    headers = ["序号", "材料类别", "材料名称", "规格型号", "单位", "单循环", "月用量", "工程总量"]
                    tbl = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
                    for j, h in enumerate(headers):
                        cell = tbl.rows[0].cells[j]
                        cell.text = h
                        for run in cell.paragraphs[0].runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for i, mat in enumerate(equipment_result.material_bom, 1):
                        row = tbl.add_row()
                        row.cells[0].text = str(i)
                        row.cells[1].text = mat.category
                        row.cells[2].text = mat.name
                        row.cells[3].text = mat.model_spec or "—"
                        row.cells[4].text = mat.unit
                        row.cells[5].text = f"{mat.qty_per_cycle}" if mat.qty_per_cycle else "—"
                        row.cells[6].text = f"{mat.qty_per_month}" if mat.qty_per_month else "—"
                        row.cells[7].text = f"{mat.qty_total}" if mat.qty_total else "—"
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(10)
                    # 计算依据
                    if equipment_result.material_bom:
                        basis = equipment_result.material_bom[0].calc_basis or ""
                        if basis:
                            p = doc.add_paragraph(f"计算依据：{basis}")
                            p.runs[0].font.size = Pt(10)
                    continue

            # 章节内容 — 智能排版引擎
            for line in ch.content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue  # 跳过空行

                # --- 优先级 1: 警告行 ---
                if stripped.startswith("⚠") or stripped.startswith("  ⚠"):
                    p = doc.add_paragraph()
                    run = p.add_run(stripped)
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    run.font.bold = True
                    continue

                # --- 优先级 2: 节标题（"第X节 XXX"） ---
                if _SECTION_TITLE_RE.match(stripped):
                    h = doc.add_heading(stripped, level=3)
                    if h.runs:
                        h.runs[0].font.size = Pt(14)
                    continue

                # 编号条目（一、二、等）用加粗段落
                if _NUMBERED_ITEM_RE.match(stripped):
                    p = doc.add_paragraph()
                    run = p.add_run(stripped)
                    run.font.size = Pt(12)
                    run.font.bold = True
                    continue

                # --- 优先级 3: 条标题（第X条）---
                if _CLAUSE_TITLE_RE.match(stripped):
                    p = doc.add_paragraph()
                    run = p.add_run(stripped)
                    run.font.size = Pt(12)
                    run.font.bold = True
                    continue

                # 数字条款：1. 2. （一）用正文缩进
                if _SUB_ITEM_RE.match(stripped):
                    p = doc.add_paragraph(stripped)
                    pf = p.paragraph_format
                    pf.left_indent = Pt(12)
                    pf.space_before = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(12)
                    continue

                # --- 优先级 4: 款标题（（1）/(1)/①等子条款） ---
                is_subclause = bool(_docx_re.match(
                    r'^(（\d+）|\(\d+\)|[①②③④⑤⑥⑦⑧⑨⑩])',
                    stripped
                ))
                if is_subclause:
                    p = doc.add_paragraph(stripped)
                    pf = p.paragraph_format
                    pf.left_indent = Pt(24)  # 缩进
                    pf.space_before = Pt(3)
                    for run in p.runs:
                        run.font.size = Pt(12)
                    continue

                # --- 优先级 5: 【】加粗标记 ---
                if stripped.startswith("【"):
                    p = doc.add_paragraph()
                    run = p.add_run(stripped)
                    run.font.bold = True
                    run.font.size = Pt(12)
                    continue

                # --- 优先级 6: 普通正文段落（首行缩进） ---
                p = doc.add_paragraph(stripped)
                pf = p.paragraph_format
                pf.first_line_indent = Pt(24)  # 首行缩进约 2 字符
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)
                pf.line_spacing = Pt(22)  # 1.5 倍行距（≈固定值 22pt）

        # --- 保存 ---
        output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
            "storage", "outputs"
        )
        os.makedirs(output_dir, exist_ok=True)

        safe_name = project.face_name.replace("/", "_").replace(" ", "_")
        filename = f"{safe_name}_作业规程_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(output_dir, filename)

        doc.save(file_path)
        return file_path
